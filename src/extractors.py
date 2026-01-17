"""
优化的文档提取器模块
专注于 PyMuPDF (fitz) 高性能提取
支持多格式并发处理

优化内容：
- PDF 提取：详细错误分类和降级策略
- 文本提取：使用 chardet 自动检测编码
- 统一错误处理和日志记录
"""

import logging
import hashlib
from pathlib import Path
from typing import Optional, Union, Iterator, List, Tuple
from dataclasses import dataclass, field
from abc import ABC, abstractmethod
import re
import enum

logger = logging.getLogger(__name__)


# ========== OCR 引擎抽象 ==========

class OCREngine(ABC):
    """OCR 引擎基类"""

    @abstractmethod
    def recognize(self, image_bytes: bytes) -> List[Tuple]:
        """
        识别图片中的文字

        Args:
            image_bytes: 图片字节数据

        Returns:
            识别结果列表，每个元素为 (bbox, text, confidence)
        """
        pass

    @abstractmethod
    def get_engine_name(self) -> str:
        """获取引擎名称"""
        pass


class RapidOCREngine(OCREngine):
    """RapidOCR 引擎（轻量级，速度快）"""

    def __init__(self):
        try:
            from rapidocr_onnxruntime import RapidOCR
            self._engine = RapidOCR()
            logger.info("RapidOCR 引擎初始化成功")
        except ImportError:
            raise ImportError("请安装 rapidocr-onnxruntime: pip install rapidocr-onnxruntime")

    def recognize(self, image_bytes: bytes) -> List[Tuple]:
        result = self._engine(image_bytes)
        if result and result[0]:
            return result[0]
        return []

    def get_engine_name(self) -> str:
        return "RapidOCR"


class PaddleOCREngine(OCREngine):
    """PaddleOCR 引擎（高精度，支持中英文混合）"""

    def __init__(self, use_gpu: bool = True, lang: str = "ch"):
        try:
            from paddleocr import PaddleOCR
            # 初始化 PaddleOCR
            self._engine = PaddleOCR(
                use_angle_cls=True,
                lang=lang,
                use_gpu=use_gpu,
                show_log=False
            )
            logger.info(f"PaddleOCR 引擎初始化成功 (lang={lang}, gpu={use_gpu})")
        except ImportError:
            raise ImportError("请安装 paddleocr: pip install paddleocr")

    def recognize(self, image_bytes: bytes) -> List[Tuple]:
        import numpy as np
        from io import BytesIO
        from PIL import Image

        # 将字节转换为 PIL Image
        img = Image.open(BytesIO(image_bytes))
        img_array = np.array(img)

        # PaddleOCR 识别
        result = self._engine.ocr(img_array, cls=True)

        # 转换为统一格式
        extracted = []
        if result and result[0]:
            for line in result[0]:
                bbox = line[0]  # 边界框
                text_info = line[1]  # (text, confidence)
                text = text_info[0]
                confidence = float(text_info[1])
                extracted.append((bbox, text, confidence))

        return extracted

    def get_engine_name(self) -> str:
        return "PaddleOCR"


class Qwen2VLOCREngine(OCREngine):
    """
    Qwen2-VL 视觉语言模型 OCR 引擎
    最强大的 OCR，支持理解复杂布局和表格
    """

    def __init__(self, model_path: str = "Qwen/Qwen2-VL-7B-Instruct", device: str = "cuda:0"):
        try:
            from transformers import Qwen2VLForConditionalGeneration, AutoProcessor
            from PIL import Image
            import torch

            logger.info(f"加载 Qwen2-VL 模型: {model_path}")

            self._model = Qwen2VLForConditionalGeneration.from_pretrained(
                model_path,
                torch_dtype=torch.bfloat16,
                device_map=device
            )
            self._processor = AutoProcessor.from_pretrained(model_path)
            self._device = device

            logger.info("Qwen2-VL 引擎初始化成功")

        except ImportError:
            raise ImportError("请安装 transformers: pip install transformers torch Pillow")

    def recognize(self, image_bytes: bytes) -> List[Tuple]:
        from io import BytesIO
        from PIL import Image

        # 将字节转换为 PIL Image
        image = Image.open(BytesIO(image_bytes))

        # 构建提示词
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "image": image,
                    },
                    {"type": "text", "text": "请识别图片中的所有文字内容，保持原有布局和格式。如果是表格，请用表格格式输出。"}
                ],
            }
        ]

        # 准备输入
        text = self._processor.apply_chat_template(
            messages, tokenize=False, add_generation_prompt=True
        )
        image_inputs, video_inputs = self._processor.process_vision_info(messages)
        inputs = self._processor(
            text=[text],
            images=image_inputs,
            videos=video_inputs,
            padding=True,
            return_tensors="pt",
        )
        inputs = inputs.to(self._device)

        # 生成
        with torch.no_grad():
            generated_ids = self._model.generate(**inputs, max_new_tokens=2048)

        generated_ids_trimmed = [
            out_ids[len(in_ids):] for in_ids, out_ids in zip(inputs.input_ids, generated_ids)
        ]
        output_text = self._processor.batch_decode(
            generated_ids_trimmed, skip_special_tokens=True, clean_up_tokenization_spaces=False
        )[0]

        # Qwen2-VL 返回完整文本，没有边界框信息
        # 返回格式: (None, text, 1.0)
        return [(None, output_text, 1.0)]

    def get_engine_name(self) -> str:
        return "Qwen2-VL"


def create_ocr_engine(engine_type: str = "rapidocr", **kwargs) -> OCREngine:
    """
    创建 OCR 引擎实例

    Args:
        engine_type: 引擎类型 ("rapidocr", "paddleocr", "qwen2_vl")
        **kwargs: 引擎特定参数

    Returns:
        OCR 引擎实例
    """
    if engine_type == "rapidocr":
        return RapidOCREngine(**kwargs)
    elif engine_type == "paddleocr":
        return PaddleOCREngine(**kwargs)
    elif engine_type == "qwen2_vl":
        return Qwen2VLOCREngine(**kwargs)
    else:
        raise ValueError(f"不支持的 OCR 引擎类型: {engine_type}")


# ========== 错误分类 ==========

class ExtractionErrorType(enum.Enum):
    """提取错误类型分类"""
    # PDF 相关错误
    PDF_ENCRYPTED = "pdf_encrypted"           # PDF 加密
    PDF_CORRUPTED = "pdf_corrupted"           # PDF 损坏
    PDF_NO_TEXT_LAYER = "pdf_no_text_layer"   # 无文本层（扫描件）
    PDF_EMPTY = "pdf_empty"                   # 空白 PDF
    PDF_PARSE_ERROR = "pdf_parse_error"       # 解析错误

    # 编码相关错误
    ENCODING_DETECTION_FAILED = "encoding_failed"  # 编码检测失败
    ENCODING_CONVERSION_ERROR = "encoding_error"   # 编码转换错误

    # 文件相关错误
    FILE_NOT_FOUND = "file_not_found"         # 文件不存在
    FILE_TOO_LARGE = "file_too_large"         # 文件过大
    FILE_PERMISSION_DENIED = "permission_denied"  # 权限拒绝

    # 依赖相关错误
    DEPENDENCY_MISSING = "dependency_missing"     # 缺少依赖
    DEPENDENCY_VERSION = "dependency_version"     # 依赖版本问题

    # 未知错误
    UNKNOWN_ERROR = "unknown_error"


@dataclass
class ExtractionError:
    """结构化的提取错误"""
    error_type: ExtractionErrorType
    message: str
    details: dict = field(default_factory=dict)
    suggestion: Optional[str] = None

    def to_dict(self) -> dict:
        """转换为字典（确保 JSON 可序列化）"""
        return {
            "type": self.error_type.value if isinstance(self.error_type, ExtractionErrorType) else str(self.error_type),
            "message": self.message,
            "details": self._sanitize_dict(self.details),
            "suggestion": self.suggestion
        }

    @staticmethod
    def _sanitize_dict(d: dict) -> dict:
        """递归清理字典，确保所有值都是 JSON 可序列化的"""
        cleaned = {}
        for k, v in d.items():
            if isinstance(v, (str, int, float, bool, type(None))):
                cleaned[k] = v
            elif isinstance(v, ExtractionErrorType):
                cleaned[k] = v.value
            elif isinstance(v, enum.Enum):
                cleaned[k] = v.value
            elif isinstance(v, dict):
                cleaned[k] = ExtractionError._sanitize_dict(v)
            elif isinstance(v, (list, tuple)):
                cleaned[k] = [ExtractionError._sanitize_value(item) for item in v]
            else:
                cleaned[k] = str(v)
        return cleaned

    @staticmethod
    def _sanitize_value(value) -> any:
        """清理单个值"""
        if isinstance(value, (str, int, float, bool, type(None))):
            return value
        elif isinstance(value, (ExtractionErrorType, enum.Enum)):
            return value.value
        elif isinstance(value, dict):
            return ExtractionError._sanitize_dict(value)
        elif isinstance(value, (list, tuple)):
            return [ExtractionError._sanitize_value(item) for item in value]
        else:
            return str(value)


# ========== 提取结果 ==========

@dataclass
class ExtractResult:
    """文档提取结果"""
    content: str           # 提取的文本内容
    metadata: dict         # 元数据
    success: bool          # 是否成功
    error: Optional[ExtractionError] = None  # 结构化错误信息

    # 兼容旧接口
    @property
    def error_message(self) -> Optional[str]:
        """获取错误消息（兼容旧接口）"""
        return self.error.message if self.error else None

    @property
    def char_count(self) -> int:
        """字符数"""
        return len(self.content)

    @property
    def word_count(self) -> int:
        """估算字数（中英文混合）"""
        # 中文计数
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', self.content))
        # 英文单词计数
        english_words = len(re.findall(r'\b[a-zA-Z]+\b', self.content))
        return chinese_chars + english_words

    @property
    def content_hash(self) -> str:
        """内容哈希（用于去重）"""
        return hashlib.md5(self.content.encode('utf-8')).hexdigest()


# ========== 基础提取器 ==========

class BaseExtractor(ABC):
    """文档提取器基类"""

    @abstractmethod
    def can_extract(self, file_path: Path) -> bool:
        """判断是否支持提取该文件"""
        pass

    @abstractmethod
    def extract(self, file_path: Path) -> ExtractResult:
        """提取文件内容"""
        pass

    def _validate_file(self, file_path: Path) -> Optional[ExtractionError]:
        """验证文件基本信息"""
        # 检查文件是否存在
        if not file_path.exists():
            return ExtractionError(
                error_type=ExtractionErrorType.FILE_NOT_FOUND,
                message=f"文件不存在: {file_path}",
                details={"path": str(file_path)},
                suggestion="请检查文件路径是否正确"
            )

        # 检查是否为文件
        if not file_path.is_file():
            return ExtractionError(
                error_type=ExtractionErrorType.FILE_NOT_FOUND,
                message=f"路径不是文件: {file_path}",
                details={"path": str(file_path)}
            )

        # 检查文件大小
        try:
            file_size = file_path.stat().st_size
            if file_size == 0:
                return ExtractionError(
                    error_type=ExtractionErrorType.FILE_TOO_LARGE,
                    message="文件为空",
                    details={"path": str(file_path)}
                )

            # 100MB 警告
            if file_size > 100 * 1024 * 1024:
                logger.warning(f"文件较大 ({file_size / 1024 / 1024:.1f}MB): {file_path.name}")

        except Exception as e:
            return ExtractionError(
                error_type=ExtractionErrorType.FILE_PERMISSION_DENIED,
                message=f"无法访问文件: {str(e)}",
                details={"path": str(file_path), "error": str(e)}
            )

        return None


# ========== PDF 提取器（优化版）==========

class FastPDFExtractor(BaseExtractor):
    """
    高性能 PDF 提取器
    使用 PyMuPDF (fitz) 优化的提取策略
    添加详细错误分类和降级策略
    支持多种 OCR 引擎的级联兜底方案

    提取策略（按优先级）：
    1. 标准文本提取（PyMuPDF get_text）
    2. 备用提取模式（多种 flags 尝试）
    3. OCR 降级级联方案：
       - RapidOCR（快速，基础识别）
       - PaddleOCR（高精度，支持表格和日文）
    """

    # OCR 引擎优先级列表（从快到强）
    OCR_ENGINE_CASCADE = [
        ("rapidocr", {"lang": "ch"}),
        ("paddleocr", {"use_gpu": True, "lang": "ch"})
    ]

    def __init__(self,
                 extract_images: bool = False,
                 extract_tables: bool = False,
                 clean_whitespace: bool = True,
                 min_line_length: int = 0,
                 enable_ocr_fallback: bool = True,
                 ocr_engine_type: str = "cascade",
                 ocr_engine_kwargs: dict = None):
        """
        初始化 PDF 提取器

        Args:
            extract_images: 是否提取图片OCR文本
            extract_tables: 是否智能提取表格
            clean_whitespace: 是否清理空白字符
            min_line_length: 最小行长度过滤
            enable_ocr_fallback: 是否启用 OCR 降级（默认启用）
            ocr_engine_type: OCR 引擎类型 ("rapidocr", "paddleocr", "cascade")
                              "cascade" 表示按优先级依次尝试所有引擎
            ocr_engine_kwargs: OCR 引擎初始化参数
        """
        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self.clean_whitespace = clean_whitespace
        self.min_line_length = min_line_length
        self.enable_ocr_fallback = enable_ocr_fallback
        self.ocr_engine_type = ocr_engine_type
        self.ocr_engine_kwargs = ocr_engine_kwargs or {}
        self._ocr_engines = {}  # 存储已初始化的 OCR 引擎

    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def extract(self, file_path: Path) -> ExtractResult:
        """提取 PDF 内容（增强版：详细错误分类和降级策略）"""
        # 验证文件
        if error := self._validate_file(file_path):
            return ExtractResult(content="", metadata={}, success=False, error=error)

        # 检查依赖
        try:
            import fitz  # PyMuPDF
        except ImportError:
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.DEPENDENCY_MISSING,
                    message="PyMuPDF 未安装",
                    suggestion="请运行: pip install PyMuPDF"
                )
            )

        # 尝试提取
        return self._extract_with_fallback(file_path, fitz)

    def _extract_with_fallback(self, file_path: Path, fitz) -> ExtractResult:
        """带降级策略的提取（支持级联 OCR）"""
        extraction_errors = []

        # 策略 1: 标准文本提取
        result = self._extract_standard(file_path, fitz)
        if result.success:
            return result
        else:
            extraction_errors.append(("standard", result.error))

        # 策略 2: 尝试不同提取模式
        result = self._extract_alternative_mode(file_path, fitz)
        if result.success:
            logger.info(f"{file_path.name}: 使用备用模式提取成功")
            return result
        else:
            extraction_errors.append(("alternative_mode", result.error))

        # 策略 3: OCR 降级级联方案（如果启用）
        if self.enable_ocr_fallback:
            if self.ocr_engine_type == "cascade":
                # 尝试所有 OCR 引擎（按优先级）
                for engine_type, default_kwargs in self.OCR_ENGINE_CASCADE:
                    result = self._extract_with_ocr(file_path, fitz, engine_type, default_kwargs)
                    if result.success:
                        logger.info(f"{file_path.name}: 使用 {engine_type} OCR 提取成功")
                        return result
                    else:
                        extraction_errors.append((f"ocr_{engine_type}", result.error))
            else:
                # 使用指定的 OCR 引擎
                result = self._extract_with_ocr(file_path, fitz, self.ocr_engine_type, self.ocr_engine_kwargs)
                if result.success:
                    logger.info(f"{file_path.name}: 使用 {self.ocr_engine_type} OCR 提取成功")
                    return result
                else:
                    extraction_errors.append(("ocr", result.error))

        # 所有策略都失败，返回详细的错误信息
        return self._create_failure_result(file_path, extraction_errors)

    def _extract_standard(self, file_path: Path, fitz) -> ExtractResult:
        """标准提取模式"""
        try:
            doc = fitz.open(file_path)

            # 检查是否加密
            if doc.is_encrypted:
                # 尝试空密码解密
                try:
                    doc.authenticate("")
                except Exception:
                    return ExtractResult(
                        content="",
                        metadata={},
                        success=False,
                        error=ExtractionError(
                            error_type=ExtractionErrorType.PDF_ENCRYPTED,
                            message="PDF 文件已加密",
                            details={"file": file_path.name},
                            suggestion="请提供解密后的 PDF 文件"
                        )
                    )

            # 基础元数据
            metadata = {
                "pages": len(doc),
                "format": "pdf",
                "encrypted": doc.is_encrypted,
                "extractor": "FastPDFExtractor.standard"
            }

            # 检查是否为空白 PDF
            if len(doc) == 0:
                return ExtractResult(
                    content="",
                    metadata=metadata,
                    success=False,
                    error=ExtractionError(
                        error_type=ExtractionErrorType.PDF_EMPTY,
                        message="PDF 文件没有页面",
                        details={"file": file_path.name}
                    )
                )

            # 提取目录元数据
            if toc := doc.get_toc():
                metadata["toc_items"] = len(toc)

            # 快速文本提取
            content_parts = []
            empty_pages = 0

            for page_num, page in enumerate(doc):
                text = self._extract_page_text(page, page_num)

                if text:
                    content_parts.append(text)
                else:
                    empty_pages += 1

            doc.close()

            # 合并内容
            content = "\n\n".join(content_parts)

            # 后处理
            if self.clean_whitespace:
                content = self._clean_text(content)

            metadata.update({
                "char_count": len(content),
                "empty_pages": empty_pages,
                "extraction_method": "standard"
            })

            # 判断是否成功
            if len(content) <= self.min_line_length:
                return ExtractResult(
                    content=content,
                    metadata=metadata,
                    success=False,
                    error=ExtractionError(
                        error_type=ExtractionErrorType.PDF_NO_TEXT_LAYER,
                        message=f"PDF 提取内容过少 ({len(content)} 字符)，可能是扫描件",
                        details={
                            "file": file_path.name,
                            "char_count": len(content),
                            "empty_pages": empty_pages,
                            "total_pages": len(doc)
                        },
                        suggestion="如果启用 OCR，可以尝试 OCR 提取"
                    )
                )

            return ExtractResult(
                content=content,
                metadata=metadata,
                success=True
            )

        except fitz.FileDataError as e:
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.PDF_CORRUPTED,
                    message=f"PDF 文件损坏: {str(e)}",
                    details={"file": file_path.name, "error": str(e)},
                    suggestion="请检查 PDF 文件是否完整"
                )
            )
        except Exception as e:
            logger.debug(f"标准提取失败 {file_path.name}: {e}")
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.PDF_PARSE_ERROR,
                    message=f"标准提取失败: {str(e)}",
                    details={"file": file_path.name, "error": str(e)}
                )
            )

    def _extract_alternative_mode(self, file_path: Path, fitz) -> ExtractResult:
        """备用提取模式（尝试不同的文本提取选项）"""
        try:
            doc = fitz.open(file_path)

            if doc.is_encrypted:
                doc.authenticate("")

            content_parts = []

            for page_num, page in enumerate(doc):
                # 尝试多种提取模式
                texts_to_try = [
                    # 保留空白和连字符
                    page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_PRESERVE_LIGATURES),
                    # 按块提取
                    "\n".join([b[4] for b in page.get_text("blocks") if b[6] == 0]),
                    # 保留布局
                    page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE),
                ]

                # 选择最长的结果
                best_text = max(texts_to_try, key=len)
                if best_text.strip():
                    content_parts.append(best_text.strip())

            doc.close()

            content = "\n\n".join(content_parts)

            if self.clean_whitespace:
                content = self._clean_text(content)

            if len(content) > self.min_line_length:
                return ExtractResult(
                    content=content,
                    metadata={
                        "pages": len(content_parts),
                        "format": "pdf",
                        "extractor": "FastPDFExtractor.alternative",
                        "char_count": len(content),
                        "extraction_method": "alternative"
                    },
                    success=True
                )

            return ExtractResult(
                content=content,
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.PDF_NO_TEXT_LAYER,
                    message="备用模式提取内容过少",
                    details={"char_count": len(content)}
                )
            )

        except Exception as e:
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.PDF_PARSE_ERROR,
                    message=f"备用提取失败: {str(e)}",
                    details={"error": str(e)}
                )
            )

    def _extract_with_ocr(self, file_path: Path, fitz, engine_type: str = None, engine_kwargs: dict = None) -> ExtractResult:
        """
        OCR 降级提取（使用指定的 OCR 引擎）

        Args:
            file_path: PDF 文件路径
            fitz: PyMuPDF 模块
            engine_type: OCR 引擎类型
            engine_kwargs: OCR 引擎参数
        """
        engine_type = engine_type or self.ocr_engine_type
        engine_kwargs = engine_kwargs or self.ocr_engine_kwargs

        # 获取或创建 OCR 引擎
        if engine_type not in self._ocr_engines:
            try:
                # 使用统一的 OCR 引擎工厂创建
                self._ocr_engines[engine_type] = create_ocr_engine(engine_type, **engine_kwargs)
                logger.info(f"OCR 引擎初始化成功: {engine_type}")
            except (ImportError, ValueError) as e:
                return ExtractResult(
                    content="",
                    metadata={},
                    success=False,
                    error=ExtractionError(
                        error_type=ExtractionErrorType.DEPENDENCY_MISSING,
                        message=f"OCR 引擎初始化失败 ({engine_type}): {str(e)}",
                        suggestion=f"请确保 {engine_type} 的依赖已安装"
                    )
                )

        ocr_engine = self._ocr_engines[engine_type]

        try:
            doc = fitz.open(file_path)
            if doc.is_encrypted:
                doc.authenticate("")

            content_parts = []
            ocr_used_pages = 0

            for page_num, page in enumerate(doc):
                # 将页面渲染为图片
                mat = fitz.Matrix(200 / 72, 200 / 72)  # 200 DPI
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                pix = None

                # OCR 识别（使用统一的引擎接口）
                result = ocr_engine.recognize(img_bytes)

                if result:
                    # result 格式: [(bbox, text, confidence), ...]
                    lines = [line[1] for line in result if line and len(line) > 1 and line[1]]
                    if lines:
                        text = "\n".join(lines)
                        content_parts.append(text)
                        ocr_used_pages += 1

            doc.close()

            content = "\n\n".join(content_parts)

            if self.clean_whitespace:
                content = self._clean_text(content)

            if len(content) > self.min_line_length:
                return ExtractResult(
                    content=content,
                    metadata={
                        "pages": len(content_parts),
                        "format": "pdf",
                        "extractor": f"FastPDFExtractor.ocr.{ocr_engine.get_engine_name()}",
                        "char_count": len(content),
                        "ocr_used_pages": ocr_used_pages,
                        "extraction_method": "ocr",
                        "ocr_engine": ocr_engine.get_engine_name()
                    },
                    success=True
                )

            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.PDF_NO_TEXT_LAYER,
                    message=f"OCR 提取内容过少 ({len(content)} 字符)",
                    details={"char_count": len(content), "engine": engine_type}
                )
            )

        except Exception as e:
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.PDF_PARSE_ERROR,
                    message=f"OCR 提取失败 ({engine_type}): {str(e)}",
                    details={"error": str(e), "engine": engine_type}
                )
            )

    def _create_failure_result(self, file_path: Path, errors: List[Tuple[str, ExtractionError]]) -> ExtractResult:
        """创建失败结果（汇总所有尝试）"""
        error_details = []
        for method, error in errors:
            error_details.append(f"{method}: {error.message}")

        return ExtractResult(
            content="",
            metadata={"file": file_path.name, "attempts": len(errors)},
            success=False,
            error=ExtractionError(
                error_type=ExtractionErrorType.UNKNOWN_ERROR,
                message="所有提取策略都失败",
                details={
                    "attempts": error_details,
                    "suggestion": "如果是扫描件，请启用 OCR；如果文件损坏，请重新获取"
                }
            )
        )

    def _extract_page_text(self, page, page_num: int) -> str:
        """提取单页文本"""
        text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        if self.extract_tables:
            try:
                tables = page.find_tables()
                if tables.tables:
                    for table in tables:
                        table_text = self._extract_table(table)
                        if table_text:
                            text += "\n\n" + table_text
            except Exception as e:
                logger.debug(f"表格提取失败 (page {page_num}): {e}")

        return text.strip()

    def _extract_table(self, table) -> str:
        """提取表格内容"""
        try:
            rows = []
            for row in table.extract():
                row_text = " | ".join([str(cell) for cell in row])
                rows.append(row_text)
            return "\n".join(rows)
        except Exception:
            return ""

    def _clean_text(self, text: str) -> str:
        """清理文本"""
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = '\n'.join(line.strip() for line in text.split('\n'))
        return text.strip()


# ========== 文本提取器（优化版：使用 chardet）==========

class FastTextExtractor(BaseExtractor):
    """高性能文本文件提取器（增强版：智能编码检测）"""

    # 常见编码列表（按优先级）
    COMMON_ENCODINGS = [
        "utf-8",
        "utf-8-sig",  # UTF-8 with BOM
        "gbk",        # 中文简体
        "gb2312",     # 中文简体（旧）
        "gb18030",    # 中文国标
        "big5",       # 中文繁体
        "shift_jis",  # 日文
        "euc_jp",     # 日文
        "euc_kr",     # 韩文
        "utf-16",
        "utf-32",
        "latin-1",    # 西欧
        "iso-8859-1",
        "cp1252",     # Windows 西欧
    ]

    def __init__(self, use_chardet: bool = True, fallback_to_ignore: bool = True):
        """
        初始化文本提取器

        Args:
            use_chardet: 是否使用 chardet 自动检测编码
            fallback_to_ignore: 是否在所有编码失败时使用 ignore 模式
        """
        self.use_chardet = use_chardet
        self.fallback_to_ignore = fallback_to_ignore
        self._chardet_available = None

    def _check_chardet(self) -> bool:
        """检查 chardet 是否可用"""
        if self._chardet_available is None:
            try:
                import chardet
                self._chardet_available = True
            except ImportError:
                self._chardet_available = False
                if self.use_chardet:
                    logger.warning("chardet 未安装，将使用预定义编码列表。建议安装: pip install chardet")
        return self._chardet_available

    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".txt", ".md", ".csv", ".json", ".xml", ".log", ".conf", ".ini"}

    def extract(self, file_path: Path) -> ExtractResult:
        """提取文本文件内容（增强版：智能编码检测）"""
        # 验证文件
        if error := self._validate_file(file_path):
            return ExtractResult(content="", metadata={}, success=False, error=error)

        extraction_attempts = []

        # 策略 1: 使用 chardet 自动检测编码
        if self.use_chardet and self._check_chardet():
            result = self._extract_with_chardet(file_path)
            extraction_attempts.append(("chardet", result))
            if result.success:
                return result

        # 策略 2: 尝试常见编码
        result = self._extract_with_common_encodings(file_path)
        extraction_attempts.append(("common_encodings", result))
        if result.success:
            return result

        # 策略 3: 使用 ignore 模式（最后手段）
        if self.fallback_to_ignore:
            result = self._extract_with_ignore(file_path)
            extraction_attempts.append(("ignore_mode", result))
            if result.success:
                logger.warning(f"{file_path.name}: 使用 ignore 模式读取，可能有字符丢失")
                return result

        # 所有策略都失败
        return ExtractResult(
            content="",
            metadata={"file": file_path.name, "attempts": len(extraction_attempts)},
            success=False,
            error=ExtractionError(
                error_type=ExtractionErrorType.ENCODING_CONVERSION_ERROR,
                message="所有编码检测策略都失败",
                details={
                    "attempts": [f"{name}: {error or 'success'}" for name, (result, error) in
                                [(name, (r, r.error.message if r.error else None)) for name, r in extraction_attempts]],
                    "suggestion": "请检查文件是否为文本文件，或尝试使用专用编辑器转换编码"
                },
                suggestion="请检查文件是否为有效的文本文件"
            )
        )

    def _extract_with_chardet(self, file_path: Path) -> ExtractResult:
        """使用 chardet 自动检测编码"""
        try:
            import chardet

            # 先读取一部分内容检测编码
            with open(file_path, "rb") as f:
                raw_data = f.read(10000)  # 读取前 10KB

            detection = chardet.detect(raw_data)
            detected_encoding = detection.get("encoding")
            confidence = detection.get("confidence", 0)

            if not detected_encoding:
                return ExtractResult(
                    content="",
                    metadata={},
                    success=False,
                    error=ExtractionError(
                        error_type=ExtractionErrorType.ENCODING_DETECTION_FAILED,
                        message="chardet 无法检测编码"
                    )
                )

            # 使用检测到的编码读取
            try:
                with open(file_path, "r", encoding=detected_encoding, errors="strict") as f:
                    content = f.read()

                metadata = {
                    "format": file_path.suffix[1:],
                    "encoding": detected_encoding,
                    "encoding_confidence": f"{confidence:.2%}",
                    "encoding_method": "chardet",
                    "char_count": len(content),
                    "extractor": "FastTextExtractor.chardet"
                }

                return ExtractResult(content=content, metadata=metadata, success=True)

            except UnicodeDecodeError as e:
                return ExtractResult(
                    content="",
                    metadata={},
                    success=False,
                    error=ExtractionError(
                        error_type=ExtractionErrorType.ENCODING_CONVERSION_ERROR,
                        message=f"chardet 检测的编码 ({detected_encoding}) 解码失败",
                        details={"detected_encoding": detected_encoding, "confidence": confidence, "error": str(e)}
                    )
                )

        except Exception as e:
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.ENCODING_DETECTION_FAILED,
                    message=f"chardet 检测失败: {str(e)}",
                    details={"error": str(e)}
                )
            )

    def _extract_with_common_encodings(self, file_path: Path) -> ExtractResult:
        """使用常见编码列表尝试"""
        last_error = None

        for encoding in self.COMMON_ENCODINGS:
            try:
                with open(file_path, "r", encoding=encoding, errors="strict") as f:
                    content = f.read()

                metadata = {
                    "format": file_path.suffix[1:],
                    "encoding": encoding,
                    "encoding_method": "common_list",
                    "char_count": len(content),
                    "extractor": "FastTextExtractor.common"
                }

                return ExtractResult(content=content, metadata=metadata, success=True)

            except (UnicodeDecodeError, UnicodeError) as e:
                last_error = e
                continue
            except Exception as e:
                last_error = e
                break

        return ExtractResult(
            content="",
            metadata={},
            success=False,
            error=ExtractionError(
                error_type=ExtractionErrorType.ENCODING_CONVERSION_ERROR,
                message=f"尝试了 {len(self.COMMON_ENCODINGS)} 种常见编码都失败",
                details={"encodings_tried": self.COMMON_ENCODINGS, "last_error": str(last_error)}
            )
        )

    def _extract_with_ignore(self, file_path: Path) -> ExtractResult:
        """使用 ignore 模式读取（会丢弃无法解码的字符）"""
        try:
            # 尝试 UTF-8 with ignore
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            if content:  # 至少读取到一些内容
                metadata = {
                    "format": file_path.suffix[1:],
                    "encoding": "utf-8 (ignore)",
                    "encoding_method": "ignore",
                    "char_count": len(content),
                    "extractor": "FastTextExtractor.ignore"
                }
                return ExtractResult(content=content, metadata=metadata, success=True)

            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.ENCODING_CONVERSION_ERROR,
                    message="ignore 模式读取后内容为空"
                )
            )

        except Exception as e:
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.ENCODING_CONVERSION_ERROR,
                    message=f"ignore 模式读取失败: {str(e)}",
                    details={"error": str(e)}
                )
            )


# ========== DOCX 提取器 ==========

class FastDocxExtractor(BaseExtractor):
    """高性能 Word 文档提取器"""

    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def extract(self, file_path: Path) -> ExtractResult:
        """提取 DOCX 内容"""
        # 验证文件
        if error := self._validate_file(file_path):
            return ExtractResult(content="", metadata={}, success=False, error=error)

        try:
            from docx import Document

            doc = Document(file_path)

            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # 提取表格
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)

            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n" + "\n".join(tables_text)

            metadata = {
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "char_count": len(content),
                "extractor": "FastDocxExtractor"
            }

            return ExtractResult(content=content, metadata=metadata, success=True)

        except ImportError:
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.DEPENDENCY_MISSING,
                    message="python-docx 未安装",
                    suggestion="请运行: pip install python-docx"
                )
            )
        except Exception as e:
            return ExtractResult(
                content="",
                metadata={},
                success=False,
                error=ExtractionError(
                    error_type=ExtractionErrorType.UNKNOWN_ERROR,
                    message=f"DOCX 提取失败: {str(e)}",
                    details={"error": str(e)}
                )
            )


# ========== 统一文档提取器 ==========

class DocumentExtractor:
    """
    高性能文档提取器
    统一接口，自动路由到最优提取器
    支持多种 OCR 引擎配置和级联兜底方案

    OCR 策略说明：
    - ocr_engine_type="cascade": 自动尝试 RapidOCR → PaddleOCR
    - ocr_engine_type="rapidocr": 仅使用 RapidOCR（快速）
    - ocr_engine_type="paddleocr": 仅使用 PaddleOCR（高精度，支持表格）
    """

    def __init__(self,
                 enable_ocr: bool = True,
                 ocr_threshold: int = 50,
                 use_chardet: bool = True,
                 ocr_engine_type: str = "cascade",
                 ocr_engine_kwargs: dict = None):
        """
        初始化提取器链

        Args:
            enable_ocr: 是否启用 OCR 支持（默认启用，强烈推荐）
            ocr_threshold: OCR 字符数阈值
            use_chardet: 是否使用 chardet 自动检测编码
            ocr_engine_type: OCR 引擎类型 ("rapidocr", "paddleocr", "cascade")
            ocr_engine_kwargs: OCR 引擎初始化参数
        """
        self.use_chardet = use_chardet
        self.ocr_engine_type = ocr_engine_type
        self.extractors: list[BaseExtractor] = [
            FastTextExtractor(use_chardet=use_chardet),
            FastDocxExtractor(),
        ]

        # PDF 提取器：根据是否启用 OCR 选择
        if enable_ocr:
            # 使用支持 OCR 降级的提取器（默认级联模式）
            self.extractors.insert(0, FastPDFExtractor(
                enable_ocr_fallback=True,
                ocr_engine_type=ocr_engine_type,
                ocr_engine_kwargs=ocr_engine_kwargs
            ))
            if ocr_engine_type == "cascade":
                logger.info(f"已启用级联 OCR 降级支持 (RapidOCR → PaddleOCR)")
            else:
                logger.info(f"已启用 OCR 降级支持 (引擎: {ocr_engine_type})")
        else:
            # 使用快速 PDF 提取器（仅文本层）
            self.extractors.insert(0, FastPDFExtractor(enable_ocr_fallback=False))
            logger.info("使用快速 PDF 提取器（无 OCR）")

        if use_chardet:
            logger.info("文本编码检测: chardet 自动检测")
        else:
            logger.info("文本编码检测: 预定义编码列表")

    def add_extractor(self, extractor: BaseExtractor) -> None:
        """添加自定义提取器"""
        self.extractors.insert(0, extractor)

    def extract(self, file_path: Union[str, Path]) -> ExtractResult:
        """
        提取文件内容

        Args:
            file_path: 文件路径

        Returns:
            提取结果
        """
        file_path = Path(file_path)

        # 查找合适的提取器
        for extractor in self.extractors:
            if extractor.can_extract(file_path):
                try:
                    result = extractor.extract(file_path)

                    # 记录提取结果
                    if result.success:
                        logger.debug(f"成功提取 {file_path.name}: {result.char_count} 字符")
                    else:
                        logger.warning(f"提取失败 {file_path.name}: {result.error.message if result.error else '未知错误'}")

                    return result

                except Exception as e:
                    logger.error(f"提取器异常 {file_path.name}: {e}")
                    return ExtractResult(
                        content="",
                        metadata={},
                        success=False,
                        error=ExtractionError(
                            error_type=ExtractionErrorType.UNKNOWN_ERROR,
                            message=f"提取器异常: {str(e)}",
                            details={"extractor": type(extractor).__name__}
                        )
                    )

        # 没有找到合适的提取器
        return ExtractResult(
            content="",
            metadata={},
            success=False,
            error=ExtractionError(
                error_type=ExtractionErrorType.UNKNOWN_ERROR,
                message=f"不支持的文件格式: {file_path.suffix}",
                details={"file": str(file_path), "suffix": file_path.suffix},
                suggestion=f"支持的格式: .pdf, .txt, .md, .csv, .json, .xml, .docx"
            )
        )

    def extract_directory(self,
                         directory: Union[str, Path],
                         recursive: bool = True,
                         min_length: int = 10) -> Iterator[tuple[Path, ExtractResult]]:
        """
        批量提取目录下的文件

        Args:
            directory: 目录路径
            recursive: 是否递归处理子目录
            min_length: 最小内容长度过滤

        Yields:
            (文件路径, 提取结果) 元组
        """
        directory = Path(directory)

        if not directory.is_dir():
            raise ValueError(f"不是目录: {directory}")

        pattern = "**/*" if recursive else "*"

        for file_path in directory.glob(pattern):
            if file_path.is_file():
                result = self.extract(file_path)

                # 过滤过短内容
                if result.success and result.char_count < min_length:
                    logger.debug(f"跳过过短文件 {file_path.name}: {result.char_count} 字符")
                    continue

                yield file_path, result


# ========== 便捷函数 ==========

def get_extractor(enable_ocr: bool = True,
                  use_chardet: bool = True,
                  ocr_engine_type: str = "cascade",
                  ocr_engine_kwargs: dict = None) -> DocumentExtractor:
    """
    获取文档提取器实例

    Args:
        enable_ocr: 是否启用 OCR 支持（默认启用）
        use_chardet: 是否使用 chardet 自动检测编码
        ocr_engine_type: OCR 引擎类型 ("rapidocr", "paddleocr", "cascade")
                        默认 "cascade" 会自动尝试所有引擎
        ocr_engine_kwargs: OCR 引擎初始化参数

    Returns:
        文档提取器实例
    """
    return DocumentExtractor(
        enable_ocr=enable_ocr,
        use_chardet=use_chardet,
        ocr_engine_type=ocr_engine_type,
        ocr_engine_kwargs=ocr_engine_kwargs
    )


# ========== 兼容旧接口 ==========

# 保留旧的简单错误类型用于兼容
class SimpleExtractResult:
    """兼容旧接口的提取结果"""
    def __init__(self, content: str, metadata: dict, success: bool, error: Optional[str] = None):
        # 将新的 ExtractResult 转换为旧格式
        if isinstance(error, ExtractionError):
            error = error.message
        self.content = content
        self.metadata = metadata
        self.success = success
        self.error = error
        self.char_count = len(content)
